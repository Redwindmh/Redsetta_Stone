import docx
import pypdf
import docx2txt
from deep_translator import (GoogleTranslator,
                             ChatGptTranslator,
                             DeeplTranslator)
from config import OPENAI_API_KEY, DEEPL_API_KEY

file_path = "to_be_translated/1.docx"
finished_path = "translated"

doc = docx.Document()

doc.add_heading('This might work', 0)

def translate_doc():
    translation = ChatGptTranslator(api_key = OPENAI_API_KEY, source = "ja", target = "en").translate_file(file_path)

    doc.add_paragraph(translation)

    doc.save(f"{finished_path}/test.docx")

translate_doc()
